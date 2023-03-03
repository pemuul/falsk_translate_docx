from deep_translator import GoogleTranslator
import docx
# pip docx2txt
# pip deep-translator

def translate_docx(source_name, target_name, source_language='russian', target_language='english'):    
    ''' переводим docx файл '''
    
    doc = docx.Document(source_name)
    paragraphs = [paragraph.text for paragraph in doc.paragraphs]

    doc_save = docx.Document()
    for i, paragraph in enumerate(paragraphs):
        try:
            translation = GoogleTranslator(source=source_language, target=target_language).translate(paragraph)
            doc_save.add_paragraph(translation)
            print("Success "+str(i))
        except:
            print("Error "+str(i))

    doc_save.save(target_name)

    return target_name